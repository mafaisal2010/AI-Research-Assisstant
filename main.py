# main.py
import streamlit as st
from research_module import get_wikipedia_summary_and_full_text, clean_wikipedia_text, get_web_summaries
from io import BytesIO
from fpdf import FPDF
from docx import Document

st.set_page_config(page_title="AI Research Assistant", layout="wide")

st.title("🔍 AI Research Assistant")
query = st.text_input("Enter your research topic:", "Artificial Intelligence")

if query:
    wiki_summary, full_text, wiki_url = get_wikipedia_summary_and_full_text(query)
    cleaned_full_text = clean_wikipedia_text(full_text, wiki_url) if full_text else ""

    st.subheader("Wikipedia Summary")
    if wiki_summary:
        st.write(wiki_summary)
        if st.button("Tell me more"):
            if cleaned_full_text:
                st.text_area("Full Wikipedia Article", cleaned_full_text, height=400)
            else:
                st.warning("Could not fetch full article.")
    else:
        st.write("No Wikipedia summary found.")

    st.subheader("Web Summaries")
    link_summaries = get_web_summaries(query)
    if link_summaries:
        for link, summary in link_summaries:
            st.markdown(f"**[{link}]({link})**")
            st.write(summary)
    else:
        st.write(" ")  # Don’t show no summaries message if wiki summary exists

    # Export section
    st.subheader("Export")
    export_format = st.selectbox("Choose export format:", ["PDF", "Word"])

    if st.button("Download Summary"):
        combined = f"Wikipedia Summary:\n{wiki_summary}\n\nFull Wikipedia Article:\n{cleaned_full_text}\n\nWeb Summaries:\n"
        for link, summary in link_summaries:
            combined += f"\n{link}\n{summary}\n"

        if export_format == "PDF":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in combined.split("\n"):
                try:
                    line = line.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(200, 10, txt=line, ln=True)
                except:
                    pdf.cell(200, 10, txt="[Unsupported characters removed]", ln=True)
            buffer = BytesIO()
            pdf.output(buffer)
            st.download_button("Download PDF", data=buffer.getvalue(), file_name="summary.pdf")

        elif export_format == "Word":
            doc = Document()
            doc.add_heading("Research Summary", 0)
            for line in combined.split("\n"):
                doc.add_paragraph(line)
            buffer = BytesIO()
            doc.save(buffer)
            st.download_button("Download Word Document", data=buffer.getvalue(), file_name="summary.docx")
